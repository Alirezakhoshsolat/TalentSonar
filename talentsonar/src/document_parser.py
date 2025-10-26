"""
Document Parser for Job Descriptions

Supports multiple file formats: PDF, TXT, Markdown, Word (DOCX)
"""

import os
from typing import Optional
import logging


class DocumentParser:
    """Parse job descriptions from various file formats."""
    
    def __init__(self):
        """Initialize the document parser."""
        self.logger = self._setup_logging()
        self.supported_formats = ['.pdf', '.txt', '.md', '.docx', '.doc']
    
    def _setup_logging(self) -> logging.Logger:
        """Setup logging for the parser."""
        logger = logging.getLogger(__name__)
        if not logger.handlers:
            handler = logging.StreamHandler()
            handler.setFormatter(logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            ))
            logger.addHandler(handler)
            logger.setLevel(logging.INFO)
        return logger
    
    def parse_file(self, file_path: str) -> str:
        """
        Parse job description from file.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            str: Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext not in self.supported_formats:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                f"Supported formats: {', '.join(self.supported_formats)}"
            )
        
        # Parse based on format
        if ext == '.txt' or ext == '.md':
            return self._parse_text(file_path)
        elif ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self._parse_word(file_path)
    
    def parse_uploaded_file(self, uploaded_file) -> str:
        """
        Parse job description from Streamlit uploaded file.
        
        Args:
            uploaded_file: Streamlit UploadedFile object
            
        Returns:
            str: Extracted text content
        """
        file_name = uploaded_file.name
        _, ext = os.path.splitext(file_name)
        ext = ext.lower()
        
        if ext not in self.supported_formats:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                f"Supported formats: {', '.join(self.supported_formats)}"
            )
        
        # Read content based on format
        if ext in ['.txt', '.md']:
            return uploaded_file.read().decode('utf-8')
        elif ext == '.pdf':
            return self._parse_pdf_bytes(uploaded_file.read())
        elif ext in ['.docx', '.doc']:
            return self._parse_word_bytes(uploaded_file.read())
    
    def _parse_text(self, file_path: str) -> str:
        """Parse plain text or markdown file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file."""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except ImportError:
            self.logger.warning("PyPDF2 not installed. Using fallback text extraction.")
            return self._parse_pdf_fallback(file_path)
    
    def _parse_pdf_bytes(self, file_bytes: bytes) -> str:
        """Parse PDF from bytes."""
        try:
            import PyPDF2
            import io
            
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except ImportError:
            self.logger.warning("PyPDF2 not installed. Cannot parse PDF.")
            return "[PDF content - install PyPDF2 to extract text]"
    
    def _parse_pdf_fallback(self, file_path: str) -> str:
        """Fallback PDF parsing (mock)."""
        return "[PDF content extraction requires PyPDF2 library]"
    
    def _parse_word(self, file_path: str) -> str:
        """Parse Word document."""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
        except ImportError:
            self.logger.warning("python-docx not installed. Using fallback.")
            return "[Word document - install python-docx to extract text]"
    
    def _parse_word_bytes(self, file_bytes: bytes) -> str:
        """Parse Word document from bytes."""
        try:
            import docx
            import io
            
            doc_file = io.BytesIO(file_bytes)
            doc = docx.Document(doc_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
        except ImportError:
            self.logger.warning("python-docx not installed. Cannot parse Word document.")
            return "[Word document - install python-docx to extract text]"


def parse_job_file(file_path: str) -> str:
    """
    Convenience function to parse a job description file.
    
    Args:
        file_path (str): Path to the job description file
        
    Returns:
        str: Extracted text content
    """
    parser = DocumentParser()
    return parser.parse_file(file_path)
